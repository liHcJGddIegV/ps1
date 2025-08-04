import os
import difflib
from docx import Document
from docx.shared import RGBColor

# Define the folder path
folder_path = "/Users/ygonzalez/Library/CloudStorage/OneDrive-InvenergyLLC/Sub Systems/Aircraft Detection Lighting System/Diversion"

# Function to extract text from a Word document
def extract_text(doc_path):
    doc = Document(doc_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Function to highlight differences inside a document
def highlight_differences(doc_path, reference_text):
    doc = Document(doc_path)
    current_text = extract_text(doc_path)
    
    # Create a diff comparison
    diff = difflib.ndiff(reference_text.splitlines(), current_text.splitlines())

    # Modify document by highlighting differences
    for para in doc.paragraphs:
        for line in diff:
            if line.startswith("+ "):  # Added text
                if line[2:] in para.text:
                    run = para.add_run(f" {line[2:]}")
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Highlight in red (added text)
            elif line.startswith("- "):  # Removed text
                if line[2:] in para.text:
                    run = para.add_run(f" {line[2:]}")
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Highlight in blue (removed text)
    
    # Save modified document
    highlighted_path = doc_path.replace(".docx", "_highlighted.docx")
    doc.save(highlighted_path)
    return highlighted_path

# Get all Word documents in the folder
docx_files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.endswith(".docx")]

if len(docx_files) < 2:
    print("Not enough documents to compare.")
else:
    reference_doc = docx_files[0]  # Use the first document as the reference
    reference_text = extract_text(reference_doc)

    print(f"Using {os.path.basename(reference_doc)} as the reference document.")

    # Compare with other documents and highlight differences
    for doc_path in docx_files[1:]:
        highlighted_doc = highlight_differences(doc_path, reference_text)
        print(f"Differences highlighted in: {highlighted_doc}")
